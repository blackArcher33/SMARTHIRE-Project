import re
import io

def extract_text_from_pdf(file):
    """
    Extract text from PDF file.
    
    Args:
        file: File object or file path
        
    Returns:
        str: Extracted text
    """
    try:
        import PyPDF2
        
        if hasattr(file, 'read'):
            pdf_reader = PyPDF2.PdfReader(file)
        else:
            with open(file, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
        
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text() + '\n'
        
        return text
    
    except ImportError:
        # Fallback if PyPDF2 not available
        return "PDF extraction requires PyPDF2. Please install: pip install PyPDF2"
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def extract_text_from_docx(file):
    """
    Extract text from DOCX file.
    
    Args:
        file: File object or file path
        
    Returns:
        str: Extracted text
    """
    try:
        import docx
        
        if hasattr(file, 'read'):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        
        return text
    
    except ImportError:
        # Fallback if python-docx not available
        return "DOCX extraction requires python-docx. Please install: pip install python-docx"
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


def clean_text(text):
    """
    Clean and normalize text.
    
    Args:
        text (str): Raw text
        
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep alphanumeric and basic punctuation
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    
    # Remove multiple periods
    text = re.sub(r'\.{2,}', '.', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text


def extract_skills_from_text(text):
    """
    Extract potential skills from text.
    
    Args:
        text (str): Text to analyze
        
    Returns:
        list: List of identified skills
    """
    # Common skill keywords
    skill_patterns = [
        r'\b(?:python|java|javascript|c\+\+|ruby|php|swift|kotlin)\b',
        r'\b(?:react|angular|vue|node\.?js|django|flask|spring)\b',
        r'\b(?:aws|azure|gcp|docker|kubernetes|jenkins)\b',
        r'\b(?:sql|nosql|mongodb|postgresql|mysql|oracle)\b',
        r'\b(?:machine learning|deep learning|ai|nlp|computer vision)\b',
        r'\b(?:agile|scrum|kanban|devops|ci/cd)\b',
    ]
    
    skills = []
    text_lower = text.lower()
    
    for pattern in skill_patterns:
        matches = re.findall(pattern, text_lower, re.IGNORECASE)
        skills.extend(matches)
    
    # Remove duplicates and return
    return list(set(skills))


def categorize_salary(salary):
    """
    Categorize salary into ranges.
    
    Args:
        salary (float): Salary amount
        
    Returns:
        str: Salary category
    """
    if salary < 50000:
        return 'Entry Level'
    elif salary < 100000:
        return 'Mid Level'
    elif salary < 150000:
        return 'Senior Level'
    else:
        return 'Executive Level'


def format_number(num):
    """
    Format number with commas.
    
    Args:
        num (int/float): Number to format
        
    Returns:
        str: Formatted number
    """
    return f"{num:,.0f}"


def calculate_match_percentage(matched_count, total_count):
    """
    Calculate match percentage.
    
    Args:
        matched_count (int): Number of matched items
        total_count (int): Total number of items
        
    Returns:
        float: Match percentage
    """
    if total_count == 0:
        return 0.0
    
    return round((matched_count / total_count) * 100, 1)
